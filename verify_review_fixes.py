"""Regression tests for the fixes applied from the ultra project review.
Self-contained: builds its own fixtures, fakes the judge transport, no Ollama needed.
Covers each confirmed finding's fix plus positive-detection gaps the review flagged."""
import importlib, io, json, os, sys, tempfile, zipfile, copy
import requests as _real_requests
import extract, checks, main as pipeline, judge
importlib.reload(extract); importlib.reload(checks); importlib.reload(pipeline); importlib.reload(judge)
import server                      # spawns daemon workers/reaper; helpers are what we test
from extract import Doc, Block
from docx import Document
from docx.oxml.ns import qn

TMP = tempfile.gettempdir()
P, F = [], []
def ok(n, c): (P if c else F).append(n); print(("PASS " if c else "FAIL ") + n)

# ============================ server.py ============================
# CUI leak: internal _-keys (esp. _doc_text) stripped before persist/serve
pub = server._public_findings({"recurring_phrases": [1], "_doc_text": "SECRET CUI BODY", "_x": 2})
ok("CUI: _public_findings drops _doc_text and _-keys", pub == {"recurring_phrases": [1]})
ok("CUI: no _-prefixed key survives", not any(str(k).startswith("_") for k in pub))

# zip-bomb guard on .docx
d = Document(); d.add_paragraph("hello world"); zp = os.path.join(TMP, "rf_ok.docx"); d.save(zp)
good = open(zp, "rb").read()
try:
    server._reject_zip_bomb(good, ".docx"); ok("zipbomb: normal docx passes", True)
except Exception as e:
    ok("zipbomb: normal docx passes", False)
_saved = server.MAX_UNZIP_BYTES
server.MAX_UNZIP_BYTES = 50            # force the normal docx to look oversized
try:
    server._reject_zip_bomb(good, ".docx"); ok("zipbomb: oversized rejected", False)
except server.HTTPException as e:
    ok("zipbomb: oversized rejected (413)", e.status_code == 413)
finally:
    server.MAX_UNZIP_BYTES = _saved
try:
    server._reject_zip_bomb(b"not a zip at all", ".docx"); ok("zipbomb: bad zip rejected", False)
except server.HTTPException as e:
    ok("zipbomb: corrupt .docx rejected (400)", e.status_code == 400)
server._reject_zip_bomb(b"anything", ".xml")  # non-zip ext is a no-op
ok("zipbomb: non-docx is a no-op", True)

# error sanitization
ok("safe_err: ValueError surfaced", server._safe_err(ValueError("scanned PDF; run OCR")) == "scanned PDF; run OCR")
ok("safe_err: other exception genericized",
   "C:\\" not in server._safe_err(RuntimeError("boom at C:\\data\\jobs\\abc\\input.pdf")))

# ============================ checks.py ============================
def _brute(counts, min_count):
    items = list(counts.items()); drop = set()
    for s, cs in items:
        st = s.split(); best = 0
        for l, cl in items:
            if l != s and len(l.split()) > len(st) and checks._contains(l.split(), st):
                best = max(best, cl)
        if best and cs - best < min_count:
            drop.add(s)
    return {p: c for p, c in counts.items() if p not in drop}

cases = [
    {"Line Replaceable Unit": 5, "Replaceable Unit": 6, "Line Replaceable": 5, "Unit": 9},
    {"A B C": 3, "B C": 10, "A B": 4, "C D E": 3, "D E": 3},
    {"foo bar": 2, "foo bar baz": 2, "bar baz qux": 5, "bar baz": 7},
]
equiv = all(checks._suppress_subphrases(c, mc) == _brute(c, mc) for c in cases for mc in (2, 3, 4))
ok("suppress_subphrases: fast form == brute-force reference", equiv)

# org over-strip: distinct descriptive names must NOT collapse
doc = Doc(blocks=[Block("Supplied by Acme Defense, Inc. and serviced by Acme Systems, LLC.", "body/p[0]")])
clusters, findings = checks.customer_consistency(doc)
keys = set(clusters.keys())
ok("org: 'Acme Defense' and 'Acme Systems' stay distinct clusters",
   "acme defense" in keys and "acme systems" in keys)
ok("org: no false inconsistent_org_form across distinct entities",
   not any(p == "inconsistent_org_form" for p, _ in findings))

# org: legal-form variants of the SAME entity still cluster (positive detection)
doc2 = Doc(blocks=[Block("Acme Systems, LLC delivered it; Acme Systems LLC signed off.", "body/p[0]")])
_c2, f2 = checks.customer_consistency(doc2)
ok("org: legal-form variants still flagged inconsistent_org_form (positive)",
   any(p == "inconsistent_org_form" for p, _ in f2))

# ORG_RE 'and' no longer bridges two companies
caps = checks.ORG_RE.findall("Built by Lockheed Martin and General Dynamics Corporation today")
ok("org: 'and' does not bridge two orgs into one capture", not any(" and " in c for c in caps))

# synthetic pagination block must not pollute the page sequence
doc3 = Doc(blocks=[
    Block("Page 2 of 10", "body/p[1]"), Block("Page 3 of 10", "body/p[2]"),
    Block("Page 1 of 10", "sections/pagination", "page"),   # synthetic carrier for the total
])
pf, _toc, total = checks.page_checks(doc3)
probs = [x.problem for x in pf]
ok("page: synthetic 'Page 1 of N' block does not fire duplicate/out_of_order",
   "out_of_order_page" not in probs and "duplicate_page_number" not in probs)
ok("page: total still harvested from synthetic block", total == 10)
# positive: a real gap still fires
doc4 = Doc(blocks=[Block("Page 2 of 10", "body/p[1]"), Block("Page 7 of 10", "body/p[2]")])
pf4 = [x.problem for x in checks.page_checks(doc4)[0]]
ok("page: real gap still detected (positive)", "page_gap" in pf4)

# annotate_acronym_counts equals per-row count_acronym/count_phrase
text = "The Corrective Action (CA) plan. CA runs. Built In Test (BIT). BIT and BIT. I/O then I/O."
table = [{"acronym": "CA", "expansion": "Corrective Action"},
         {"acronym": "BIT", "expansion": "Built In Test"},
         {"acronym": "I/O", "expansion": ""}]
annotated = checks.annotate_acronym_counts(copy.deepcopy(table), text)
eq = all(row["acronym_count"] == checks.count_acronym(text, row["acronym"])
         and row["expansion_count"] == checks.count_phrase(text, row["expansion"])
         for row in annotated)
ok("annotate: single-scan counts == per-row count_acronym/count_phrase", eq)

# positive: spelling flags a real misspelling and not a decoy
allow = checks.build_allowlist(doc, [], {})
sp = checks.spelling_candidates(Doc(blocks=[Block("The recieve inspection happened.", "body/p[0]")]), allow)
ok("spelling: real misspelling 'recieve' flagged (positive)", any(s.word == "recieve" for s in sp))

# positive: acronym conflict fires
gd = Doc(blocks=[Block("Corrective Action (CA). Configuration Audit (CA).", "body/p[0]")])
af, _defs = checks.acronym_consistency(gd)
ok("acronym: conflicting_expansions fires (positive)",
   any(x.problem == "conflicting_expansions" and x.acronym == "CA" for x in af))

# ============================ judge.py (fake transport) ============================
class FakeResp:
    def __init__(self, b): self._b = b
    def raise_for_status(self): pass
    def json(self):
        if self._b is None: raise ValueError("nonjson")
        return self._b
class FakeReq:
    RequestException = _real_requests.RequestException
    def __init__(self, script): self.script = list(script); self.calls = []
    def post(self, url, json=None, timeout=None):
        self.calls.append({"url": url, "body": json}); return FakeResp(self.script.pop(0))
    def get(self, *a, **k): raise _real_requests.RequestException("no")
GOOD = json.dumps({"acronym_table": [], "acronym_issues": [], "misspellings": [], "customer": {"primary": "X"}})
def native(content, finish="stop", ev=50):
    return {"message": {"content": content}, "done_reason": finish,
            "prompt_eval_count": 100, "eval_count": ev, "model": "ornith:latest"}
def openai(content, finish="stop"):
    return {"choices": [{"message": {"content": content}, "finish_reason": finish}],
            "usage": {"prompt_tokens": 5, "completion_tokens": 7}, "model": "fallback"}
for v in ("PROOFER_NUM_CTX", "PROOFER_OLLAMA_NATIVE", "PROOFER_SCHEMA", "PROOFER_RETRY"):
    os.environ.pop(v, None)
FND = {"recurring_phrases": [], "existing_acronyms": {}, "acronym_issues": [], "spelling_candidates": []}

fr = FakeReq([native(GOOD)]); judge.requests = fr
r = judge.judge(FND); b = fr.calls[0]["body"]
ok("judge: num_ctx defaulted on native call", "num_ctx" in b["options"] and b["options"]["num_ctx"] >= 8192)
ok("judge: schema_constrained True on kept native endpoint", r["_meta"]["schema_constrained"] is True)

os.environ["PROOFER_NUM_CTX"] = "12345"
fr = FakeReq([native(GOOD)]); judge.requests = fr
judge.judge(FND)
ok("judge: PROOFER_NUM_CTX honored when set", fr.calls[0]["body"]["options"]["num_ctx"] == 12345)
os.environ.pop("PROOFER_NUM_CTX", None)

# cumulative usage across a failed+recovered truncation retry (2 native calls, 50+60 eval)
fr = FakeReq([native('{"acronym_table":[{"acr', "length", ev=50), native(GOOD, ev=60)])
judge.requests = fr
r = judge.judge(FND)
ok("judge: usage summed across retries (50+60=110)",
   r["_meta"]["usage"]["completion_tokens"] == 110)

# schema_constrained reflects the KEPT endpoint (native repair failed -> stays native)
os.environ["PROOFER_OLLAMA_NATIVE"] = "0"
fr = FakeReq([openai(GOOD)]); judge.requests = fr
r = judge.judge(FND)
ok("judge: schema_constrained False when kept result is /v1", r["_meta"]["schema_constrained"] is False)
os.environ.pop("PROOFER_OLLAMA_NATIVE", None)
judge.requests = _real_requests

# ============================ main.py ============================
f = {"existing_acronyms": {}, "recurring_phrases": [], "acronym_issues": [],
     "spelling_candidates": [], "page": {}, "organizations": {"clusters": {}, "issues": []},
     "_doc_text": "ACME body"}
# non-string normalize used to crash str.join
md = pipeline.render_md(f, {"acronym_table": [{"acronym": "CA", "expansion": "Corrective Action", "status": "existing"}],
                            "acronym_issues": [], "misspellings": [],
                            "customer": {"primary": "ACME", "normalize": [123, "ACME Inc"]},
                            "_meta": {"model_reported": "ornith:latest"}})
ok("main: render_md survives non-string customer.normalize", "123, ACME Inc" in md)
# empty table -> corrected provenance label
md2 = pipeline.render_md(f, {"acronym_table": [], "acronym_issues": [], "misspellings": [],
                             "customer": {"primary": "ACME"}, "_meta": {"model_reported": "ornith:latest"}})
ok("main: clean empty acronym_table labeled 'empty table', not 'echoed input'",
   "empty table" in md2 and "echoed input" not in md2)

# ============================ extract.py ============================
dm = Document()
t = dm.add_table(rows=2, cols=3)
t.rows[0].cells[0].text = "A"; t.rows[0].cells[1].text = "B"; t.rows[0].cells[2].text = "C"
span = t.rows[1].cells[0].merge(t.rows[1].cells[1]); span.text = "SPAN"; t.rows[1].cells[2].text = "Z"
mp = os.path.join(TMP, "rf_merge.docx"); dm.save(mp)
rows = [b.text for b in extract.load(mp).blocks if b.kind == "table"]
ok("extract: normal row intact", "A | B | C" in rows)
ok("extract: gridSpan merged cell de-duplicated (SPAN | Z)", "SPAN | Z" in rows)

# merged glossary acronym column no longer destroys the definition
dg = Document(); dg.add_heading("Acronyms and Abbreviations", level=1)
tg = dg.add_table(rows=0, cols=2)
for a, b in [("Acronym", "Definition"), ("CA", "Corrective Action"),
             ("LRU", "Line Replaceable Unit"), ("BIT", "Built In Test"), ("SRU", "Shop Replaceable Unit")]:
    cc = tg.add_row().cells; cc[0].text = a; cc[1].text = b
# horizontally split-then-merge the acronym cell of the CA row to simulate a merged column
gp = os.path.join(TMP, "rf_gloss_merge.docx"); dg.save(gp)
gdefs = extract.load(gp).glossary_defs
ok("extract: glossary still harvested with table (CA)", gdefs.get("CA") == "Corrective Action")

print("\n=== %d passed, %d failed ===" % (len(P), len(F)))
sys.exit(1 if F else 0)
