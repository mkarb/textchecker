"""Document extraction: .docx + generic technical-manual XML -> normalized Doc.

The rest of the pipeline only ever sees a Doc (a list of text Blocks with a
location path), so the file format is irrelevant downstream.
"""
from __future__ import annotations
from dataclasses import dataclass, field
from pathlib import Path
import re

from lxml import etree


@dataclass
class Block:
    text: str
    path: str                 # location label, e.g. "chapter/para[2]" or "body/p[14]"
    kind: str = "body"        # body | heading | table | toc | page
    page: str | None = None   # explicit page-marker text if this block IS a page marker


@dataclass
class Doc:
    blocks: list = field(default_factory=list)            # list[Block], in document order
    source: str = ""
    glossary_defs: dict = field(default_factory=dict)     # acronym -> expansion (from schema tags)

    @property
    def text(self) -> str:
        return "\n".join(b.text for b in self.blocks)


def _ws(s):
    return re.sub(r"\s+", " ", (s or "")).strip()


def _hardened_xml_parser():
    """A fresh lxml parser locked down for UNTRUSTED input: no external entity
    resolution, no DTD load, no network, no unbounded tree -- the same threat model for
    a generic .xml upload and for the XML parts inside a .docx (footnotes/endnotes).
    A new instance per call because lxml parsers are not safe to share across threads
    (Stage A parses on a thread pool)."""
    return etree.XMLParser(
        recover=True, remove_comments=True, remove_pis=True,
        resolve_entities=False, load_dtd=False, no_network=True, huge_tree=False,
    )


# --------------------------------------------------------------------------- docx
# ------------------------------------------------- docx coverage: sections & parts
# python-docx's body walk misses four whole classes of content:
#   headers/footers (separate section parts -- where "Page X of Y" lives),
#   footnotes/endnotes (separate package parts), text boxes (inside drawings),
#   and tables nested in table cells. The helpers below recover all four.
_WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_PAGE_OF_RE = re.compile(r"\bPage\s+(\d+)\s+of\s+(\d+)\b", re.I)


def _w_text(el):
    """All literal text under el: every w:t, which includes cached field results
    (even inside w:fldSimple, which python-docx's .text silently drops) and never
    includes w:instrText field instructions."""
    return _ws("".join(t.text or "" for t in el.iter(_WNS + "t")))


def _docx_hf_blocks(d):
    """Header/footer blocks + the page total. Footers are per-SECTION templates --
    one footer renders on every page of its section -- so a 'Page X of Y' found here
    must NOT be emitted as a page-sequence marker (two sections' cached values would
    fake a page_gap). Instead: harvest Y into the total, strip the marker from the
    emitted text, and let the caller synthesize a single pagination block."""
    blocks, totals, seen = [], set(), set()
    for si, sec in enumerate(d.sections):
        for name in ("header", "footer", "first_page_header", "first_page_footer",
                     "even_page_header", "even_page_footer"):
            hf = getattr(sec, name, None)
            if hf is None or hf.is_linked_to_previous:
                continue                       # linked = inherits previous section's part
            pid = id(hf._element)
            if pid in seen:
                continue
            seen.add(pid)
            short = "header" if "header" in name else "footer"
            items = [(p._p, f"section[{si}]/{short}/p[{j}]")
                     for j, p in enumerate(hf.paragraphs)]
            for ti, tbl in enumerate(getattr(hf, "tables", []) or []):
                for ri, row in enumerate(tbl.rows):
                    items.append((row._tr, f"section[{si}]/{short}/tbl[{ti}]/tr[{ri}]"))
            for el, path in items:
                txt = _w_text(el)
                if not txt:
                    continue
                for m in _PAGE_OF_RE.finditer(txt):
                    totals.add(int(m.group(2)))
                residual = _ws(_PAGE_OF_RE.sub(" ", txt))
                if residual:
                    blocks.append(Block(residual, path, short))
    return blocks, (max(totals) if totals else None)


def _docx_note_blocks(d):
    """Footnotes/endnotes from their package parts (python-docx has no API for them)."""
    blocks = []
    for reltail, tag in (("/footnotes", "footnote"), ("/endnotes", "endnote")):
        part = None
        for rel in d.part.rels.values():
            if rel.reltype.endswith(reltail):
                part = rel.target_part
                break
        if part is None:
            continue
        try:
            root = etree.fromstring(part.blob, _hardened_xml_parser())
        except Exception:
            continue
        for note in root.iter(_WNS + tag):
            if note.get(_WNS + "type") in ("separator", "continuationSeparator"):
                continue
            txt = _w_text(note)
            if txt:
                blocks.append(Block(txt, f"{tag}s/{tag}[{note.get(_WNS + 'id')}]", tag))
    return blocks


def _docx_textbox_blocks(d):
    """Text inside shapes/text boxes (w:txbxContent). Word writes each box twice --
    mc:Choice (modern) and mc:Fallback (legacy VML) -- so Fallback copies are skipped
    to avoid double extraction."""
    blocks, n = [], 0
    for tx in d.element.body.iter(_WNS + "txbxContent"):
        anc, skip = tx.getparent(), False
        while anc is not None:
            if etree.QName(anc).localname == "Fallback":
                skip = True
                break
            anc = anc.getparent()
        if skip:
            continue
        txt = _w_text(tx)
        if txt:
            blocks.append(Block(txt, f"body/textbox[{n}]", "textbox"))
            n += 1
    return blocks


def _walk_docx_table(table, base, blocks):
    """Emit a table's rows as blocks and recurse into tables nested in cells.

    Deduplicate by underlying <w:tc> identity: python-docx's row.cells yields the SAME
    _Cell for every grid column a horizontal gridSpan covers AND for every row a vertical
    vMerge continues, so without this a merged cell's text is repeated -- 'SPAN | SPAN | Z',
    or a merged glossary acronym column flattening to 'CA | CA | Corrective Action', which
    _gloss_pair then drops as exp==acr, silently destroying the definition. cell.text does
    not include nested-table text, so nested tables are recursed explicitly."""
    # Hold the <w:tc> elements and compare by identity (`is`), NOT id(): lxml element
    # proxies are ephemeral, so a set of id()s recycles freed addresses and false-matches
    # unrelated cells. python-docx returns the SAME _tc for every column a gridSpan covers
    # and every row a vMerge continues, so identity is exactly the merge test.
    seen_tc = []                          # table-wide: catches gridSpan (in-row) and vMerge (cross-row)
    for r, row in enumerate(table.rows):
        cells, nested_here = [], []
        for c in row.cells:
            tc = c._tc
            if any(tc is s for s in seen_tc):
                continue
            seen_tc.append(tc)
            t = _ws(c.text)
            if t:
                cells.append(t)
            nested_here.extend(c.tables)
        txt = " | ".join(cells)
        if txt:
            blocks.append(Block(txt, f"{base}/tr[{r}]", "table"))
        for k, nested in enumerate(nested_here):
            _walk_docx_table(nested, f"{base}/tr[{r}]/tbl[{k}]", blocks)


def from_docx(path) -> Doc:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    d = Document(str(path))
    doc = Doc(source=str(path))
    body = d.element.body
    idx = 0
    for child in body.iterchildren():
        tag = etree.QName(child).localname
        if tag == "p":
            para = Paragraph(child, d)
            txt = _ws(para.text)
            if txt:
                style = (para.style.name if para.style is not None else "") or ""
                kind = "heading" if style.lower().startswith("heading") else "body"
                if txt.lower() in ("table of contents", "contents"):
                    kind = "toc"
                doc.blocks.append(Block(txt, f"body/p[{idx}]", kind))
            idx += 1
        elif tag == "tbl":
            _walk_docx_table(Table(child, d), f"body/tbl[{idx}]", doc.blocks)
            idx += 1
    doc.blocks.extend(_docx_textbox_blocks(d))
    doc.blocks.extend(_docx_note_blocks(d))
    hf_blocks, page_total = _docx_hf_blocks(d)
    doc.blocks.extend(hf_blocks)
    if page_total:
        # one synthetic marker: "pagination exists, total = N". A single block keeps
        # the page-sequence checks quiet while lighting up the XREF/TOC range checks.
        doc.blocks.append(Block(f"Page 1 of {page_total}", "sections/pagination", "page"))
    harvest_glossary(doc)
    return doc


# ----------------------------------------------------------------------- generic XML
# Classification is structural, not denylist-based:
#   - an element is INLINE (folded into running text) if its tag is known-inline,
#     OR it is an unknown tag with no element children of its own;
#   - an element is STRUCTURAL (its own block) if it has element children, or its
#     tag is a known block tag (so adjacent <para>/<li>/etc. stay separate).
# This keeps a parent's own text under mixed content and folds unknown inline
# tags, instead of silently dropping surrounding text.
INLINE_TAGS = {
    "b", "i", "em", "emphasis", "strong", "bold", "italic", "u", "underline",
    "sub", "sup", "span", "link", "xref", "ref", "cite", "acronym", "abbrev",
    "abbreviation", "term", "code", "tt", "kbd", "var", "q", "quote", "a",
    "fn", "footnote", "noteref", "symbol", "verbatimtext", "applicref",
    "changeinline", "hotspot", "uicontrol", "menucascade", "wintitle", "ph",
}
HEADING_TAGS = {"title", "heading", "head", "sectiontitle", "chaptertitle", "subtitle", "caption"}
PAGE_TAGS = {"pagenum", "pageno", "page", "folio", "pagebreak"}
TOC_TAGS = {"toc", "contents", "tableofcontents"}
GLOSSARY_PAIR_TAGS = {"glossentry", "glossitem", "abbreviationdef", "acronymdef"}
SKIP_TAGS = {"meta", "metadata", "style", "comment", "binary", "graphic", "image", "img"}
# Known block-level leaves -- elements that carry running text but should each be
# their own block even when they have no element children. Extend per schema.
BLOCK_TAGS = ({
    "para", "p", "listitem", "item", "li", "step", "substep", "stepexpansion",
    "note", "notes", "warning", "caution", "attention", "hazard", "danger", "safety",
    "table", "informaltable", "tgroup", "thead", "tbody", "row", "tr", "entry",
    "td", "th", "cell", "section", "chapter", "topic", "subsection", "part", "body",
    "frontmatter", "abstract", "summary", "preface", "procedure", "proceduralstep",
    "figure", "legend", "example", "equation", "formula", "def", "definition",
    "deflist", "termentry", "desc", "shortdesc", "prereq", "context", "result",
    "postreq", "troubleshooting", "cmd", "info", "verbatim", "screen",
    "programlisting", "codeblock", "literallayout", "blockquote", "epigraph",
    "label", "line", "listoffigures",
} | HEADING_TAGS | PAGE_TAGS | TOC_TAGS | GLOSSARY_PAIR_TAGS)


def _local(el):
    return etree.QName(el).localname.lower()


def _itertext(el):
    return _ws("".join(el.itertext()))


def _is_structural(c):
    """Should child `c` be its own block (vs folded into running inline text)?"""
    cl = _local(c)
    if cl in INLINE_TAGS:
        return False
    if any(isinstance(g.tag, str) for g in c):     # has element children -> sub-container
        return True
    return cl in BLOCK_TAGS                          # known block leaf


def _kind_of(tag):
    if tag in HEADING_TAGS:
        return "heading"
    if tag in PAGE_TAGS:
        return "page"
    if tag in TOC_TAGS:
        return "toc"
    return "body"



def from_xml(path) -> Doc:
    tree = etree.parse(str(path), _hardened_xml_parser())
    root = tree.getroot()
    doc = Doc(source=str(path))

    def add(text, el):
        t = _ws(text)
        if not t:
            return
        try:
            epath = tree.getelementpath(el)
        except Exception:
            epath = _local(el)
        kind = _kind_of(_local(el))
        doc.blocks.append(Block(t, epath, kind, page=t if kind == "page" else None))

    def walk(el):
        tag = _local(el)
        if tag in SKIP_TAGS:
            return
        raw = [c for c in el if isinstance(c.tag, str)]
        # leaf: no structural (non-skip) child -> the whole element is one block,
        # with any inline children folded into its text.
        if not any(_local(c) not in SKIP_TAGS and _is_structural(c) for c in raw):
            add(_itertext(el), el)
            return
        # container with possibly mixed content: keep the element's own loose text
        # (before/after/between structural children), recurse into structural ones.
        buf = [el.text or ""]
        for c in raw:
            cl = _local(c)
            if cl in SKIP_TAGS:
                buf.append(c.tail or "")            # preserve text after skipped element
                continue
            if _is_structural(c):
                add("".join(buf), el)               # flush loose inline text as a block
                buf = []
                walk(c)
                buf.append(c.tail or "")
            else:                                    # inline child -> fold text + tail
                buf.append("".join(c.itertext()))
                buf.append(c.tail or "")
        add("".join(buf), el)

    walk(root)

    # If the schema carries explicit acronym/glossary pairs, harvest them directly
    # (high precision -- no regex guessing needed).
    for el in root.iter():
        if not isinstance(el.tag, str):
            continue
        if _local(el) in GLOSSARY_PAIR_TAGS:
            short, long_ = None, None
            for c in el.iter():
                cl = _local(c)
                t = _itertext(c)
                if not t:
                    continue
                if cl in ("glossterm", "abbrev", "abbreviation", "acronym", "short") \
                        and re.fullmatch(r"[A-Z][A-Za-z0-9&./-]{1,9}", t):
                    short = t
                elif cl in ("glossdef", "expansion", "definition", "long", "fullform", "surfaceform"):
                    long_ = t
            if short and long_:
                doc.glossary_defs[short] = long_
    harvest_glossary(doc)
    return doc


# --------------------------------------------------------------------------- pdf
# Local, offline PDF extraction (no cloud OCR -- CUI never leaves the box). Uses
# pdfplumber (pdfminer.six + Pillow underneath; MIT/BSD-licensed, unlike AGPL PyMuPDF).
# PDF is the one format where the page-number checks are *real*: every block knows the
# physical page it sits on. The hard parts of PDF are handled here so downstream sees
# the same clean Doc as docx/xml:
#   - words are regrouped into lines, then reflowed into paragraphs (so a wrapped phrase
#     like "Line Replaceable\nUnit" stays one block and recurring-phrase detection works);
#   - end-of-line hyphenation is rejoined ("Correc-\ntive" -> "Corrective") so it doesn't
#     masquerade as a misspelling;
#   - running headers/footers that repeat across pages are dropped (a 50-page manual's
#     running title would otherwise dominate the recurring-phrase list) -- but a varying
#     "Page X of Y" footer is preserved so the page checks can use it;
#   - tables are emitted as "cell | cell" rows with kind="table", exactly like docx, so
#     the glossary harvester and table handling work unchanged.
_PDF_PAGE_OF = re.compile(r"\bPage\s+\d+\s+of\s+\d+\b", re.I)
_PDF_NUMBERED_HEADING = re.compile(r"^\d+(?:\.\d+){0,4}\.?\s+\S")


def _pdf_norm(s):
    """Normalize a line for cross-page boilerplate comparison. Digits are kept, so a
    varying 'Page 1 of 2' footer is NOT mistaken for a static running header."""
    return re.sub(r"\s+", " ", (s or "")).strip().lower()


def _pdf_size(w):
    return float(w.get("size") or w.get("height") or 0.0)


def _pdf_lines(words, ytol=3.0):
    """Group words (dicts with x0,x1,top,bottom,size,text) into lines, ordered top-to-
    bottom then left-to-right. Returns [(text, max_font_size, top)]."""
    ws = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
    lines, cur, cur_top = [], [], None
    for w in ws:
        if cur and abs(w["top"] - cur_top) > ytol:
            lines.append(cur); cur = []
        if not cur:
            cur_top = w["top"]
        cur.append(w)
    if cur:
        lines.append(cur)
    out = []
    for ln in lines:
        ln.sort(key=lambda w: w["x0"])
        text = _ws(" ".join(w["text"] for w in ln))
        if text:
            out.append((text, max((_pdf_size(w) for w in ln), default=0.0),
                        min(w["top"] for w in ln)))
    return out


def _pdf_body_size(pages_lines):
    sizes = sorted(s for pg in pages_lines for (_t, s, _y) in pg if s)
    return sizes[len(sizes) // 2] if sizes else 0.0


def _pdf_boilerplate(pages_lines, page_heights, band=0.12, frac=0.6, min_pages=4):
    """Normalized line texts that recur in the top/bottom band of a majority of pages
    (running headers/footers). Conservative: only engages at >=min_pages, and never
    strips a 'Page X of Y' line."""
    n = len(pages_lines)
    if n < min_pages:
        return set()
    from collections import Counter
    seen = Counter()
    for lines, h in zip(pages_lines, page_heights):
        band_lines = {_pdf_norm(t) for (t, _s, top) in lines
                      if top <= h * band or top >= h * (1 - band)}
        seen.update(band_lines)
    need = max(min_pages, int(frac * n))
    return {t for t, c in seen.items() if c >= need and not _PDF_PAGE_OF.search(t)}


def _pdf_is_heading(text, size, body_size):
    if len(text.split()) > 14:
        return False
    if body_size and size >= body_size * 1.15:
        return True
    if _PDF_NUMBERED_HEADING.match(text) and len(text) <= 90:
        return True
    if text.isupper() and 2 <= len(text) <= 60 and any(c.isalpha() for c in text):
        return True
    return False


def _pdf_reflow(lines, body_size, line_gap=1.7):
    """Reflow lines into (text, kind) paragraphs: break on a vertical gap, on a heading,
    or right after a heading. De-hyphenate at line joins."""
    paras, buf, buf_head, prev_top = [], [], False, None
    lh = (body_size or 10.0) * 1.2

    def flush():
        nonlocal buf, buf_head
        if buf:
            text = ""
            for piece in buf:
                if not text:
                    text = piece
                elif text.endswith("-") and piece[:1].islower():
                    text = text[:-1] + piece
                else:
                    text = text + " " + piece
            text = _ws(text)
            if text:
                paras.append((text, "heading" if buf_head else "body"))
        buf, buf_head = [], False

    for (text, size, top) in lines:
        heading = _pdf_is_heading(text, size, body_size)
        big_gap = prev_top is not None and (top - prev_top) > lh * line_gap
        if heading or big_gap or buf_head:
            flush()
        buf.append(text)
        if heading:
            buf_head = True
            flush()
        prev_top = top
    flush()
    return paras


def from_pdf(path) -> Doc:
    import pdfplumber

    doc = Doc(source=str(path))
    pages_lines, page_heights, table_blocks = [], [], []
    with pdfplumber.open(str(path)) as pdf:
        for pageno, page in enumerate(pdf.pages, start=1):
            page_heights.append(float(page.height or 792))
            bboxes = []
            try:
                found = page.find_tables()
            except Exception:
                found = []
            for ti, tbl in enumerate(found):
                try:
                    bboxes.append(tbl.bbox)
                    for r, row in enumerate(tbl.extract() or []):
                        cells = [_ws(c) for c in row if c and _ws(c)]
                        txt = " | ".join(cells)
                        if txt:
                            table_blocks.append(
                                Block(txt, f"page[{pageno}]/tbl[{ti}]/tr[{r}]", "table"))
                except Exception:
                    continue
            try:
                words = page.extract_words(extra_attrs=["size"]) or []
            except Exception:
                words = []

            def _in_table(w, _b=bboxes):
                for (x0, top, x1, bottom) in _b:
                    if (w["x0"] >= x0 - 1 and w["x1"] <= x1 + 1
                            and w["top"] >= top - 1 and w["bottom"] <= bottom + 1):
                        return True
                return False

            pages_lines.append(_pdf_lines([w for w in words if not _in_table(w)]))

    body_size = _pdf_body_size(pages_lines)
    boiler = _pdf_boilerplate(pages_lines, page_heights)
    saw_page_of = any(_PDF_PAGE_OF.search(t) for pg in pages_lines for (t, _s, _y) in pg)

    tbl_by_page = {}
    for b in table_blocks:
        pno = int(b.path.split("[")[1].split("]")[0])
        tbl_by_page.setdefault(pno, []).append(b)

    for pageno, lines in enumerate(pages_lines, start=1):
        kept = [(t, s, y) for (t, s, y) in lines
                if _pdf_norm(t) not in boiler or _PDF_PAGE_OF.search(t)]
        for idx, (text, kind) in enumerate(_pdf_reflow(kept, body_size)):
            if text.lower() in ("table of contents", "contents"):
                kind = "toc"
            doc.blocks.append(Block(text, f"page[{pageno}]/p[{idx}]", kind))
        for b in tbl_by_page.get(pageno, []):
            doc.blocks.append(b)
        if not saw_page_of:
            doc.blocks.append(Block(str(pageno), f"page[{pageno}]/folio", "page",
                                    page=str(pageno)))

    if "harvest_glossary" in globals():
        harvest_glossary(doc)

    if not any(b.text.strip() for b in doc.blocks if b.kind != "page"):
        raise ValueError(
            f"No extractable text in {path}. It's likely a scanned/image-only PDF; run a "
            f"local OCR pass first (e.g. ocrmypdf with Tesseract -- keep it on-box for CUI), "
            f"then re-run on the OCR'd PDF.")
    return doc


# --------------------------------------------------------------- glossary harvest
# An acronym/abbreviation glossary is ground truth: if a document formally lists
# "CA = Corrective Action", CA must be treated as an EXISTING acronym with exactly
# that expansion -- never re-proposed, never re-guessed. The parsers above leave such
# a table sitting in the block stream (docx flattens a row to "CA | Corrective Action";
# a tabbed list collapses to "CA Corrective Action"), so we recover the pairs here,
# format-agnostically. Precision over recall: a row must *look* like an acronym
# definition, and we only trust a run of them (ordinary prose won't produce a run).
_GLOSS_TITLE = re.compile(
    r"^(?:list of\s+|table of\s+)?"
    r"(?:acronyms?|abbreviations?|nomenclature|glossary|definitions?|"
    r"acronyms?\s*(?:,|&|and)\s*abbreviations?|abbreviations?\s*(?:,|&|and)\s*acronyms?|"
    r"acronyms?\s*(?:,|&|and)\s*definitions?)"
    r"(?:\s*(?:,|&|and)\s*(?:abbreviations?|acronyms?|definitions?|terms?))*\s*:?\s*$",
    re.I,
)
# Acronym cell: starts with a letter, no internal spaces, short, letters/digits + a few
# joiners (allows SD-18, I/O, MIL-STD-ish, FMECA). 2-12 chars after trimming.
_GLOSS_ACR = re.compile(r"^[A-Za-z][A-Za-z0-9&./\-]{1,11}$")
# Header-row / column-label words that must never be stored as a definition.
_GLOSS_HEADER = {
    "acronym", "acronyms", "abbreviation", "abbreviations", "abbr", "term", "terms",
    "symbol", "symbols", "definition", "definitions", "description", "meaning",
    "expansion", "expanded", "full form", "stands for", "title", "acronym/abbreviation",
    "item", "no", "no.", "ref", "page",
}


def _looks_like_acronym(acr):
    """Discriminate a true acronym cell from an ordinary word. Accept if it carries a
    non-letter (digit/slash/hyphen, e.g. SD-18, I/O) or is majority-uppercase (CA, FMECA,
    STORM). Reject normal words like 'Failure' or 'Derating' -- this is what stops a prose
    paragraph whose first token happens to match the shape from being read as a glossary
    entry."""
    if not _GLOSS_ACR.match(acr):
        return False
    if acr.lower() in _GLOSS_HEADER:
        return False
    letters = [c for c in acr if c.isalpha()]
    if any(not c.isalpha() for c in acr):          # has a digit / slash / hyphen
        return True
    upper = [c for c in letters if c.isupper()]
    return bool(letters) and len(upper) * 2 >= len(letters)


def _gloss_pair(text):
    """Parse one block as a glossary entry -> (acronym, expansion) or None. Accepts the
    two shapes extraction produces: a flattened table row 'ACR | Expansion' and a
    tab/space list 'ACR Expansion'."""
    t = (text or "").strip()
    if not t:
        return None
    if " | " in t:                                  # flattened table row
        cells = [c.strip() for c in t.split(" | ") if c.strip()]
        if len(cells) < 2:
            return None
        acr, exp = cells[0], cells[1]
    else:                                           # tab/space list row
        parts = t.split(None, 1)
        if len(parts) < 2:
            return None
        acr, exp = parts[0], parts[1]
    acr = acr.strip().rstrip(".:")
    exp = re.sub(r"^[\-\u2013\u2014:\s]+", "", exp.strip())   # drop a leading dash/colon
    if not _looks_like_acronym(acr):
        return None
    if exp.lower() in _GLOSS_HEADER or exp == acr or len(exp) < 3:
        return None
    if " " not in exp and len(exp) <= len(acr) + 2:  # single short token isn't an expansion
        return None
    if len(acr) >= len(exp):
        return None
    return acr, exp


def harvest_glossary(doc, min_run=4):
    """Populate doc.glossary_defs from a glossary table/list embedded in the blocks.
    Only fills keys not already set (explicit schema tags win). Detection trusts the
    longest run of consecutive glossary-like blocks: a real table (all 'table' blocks)
    at >=3 rows, a heading-introduced run at >=2, otherwise >=min_run (so loose prose,
    which never strings 4 acronym-led lines together, can't trigger it)."""
    blocks = doc.blocks
    n = len(blocks)
    pairs = [_gloss_pair(b.text) for b in blocks]
    heading_at = [bool(_GLOSS_TITLE.match((b.text or "").strip())) for b in blocks]

    runs, i = [], 0
    while i < n:
        if pairs[i]:
            j = i
            while j < n and pairs[j]:
                j += 1
            runs.append((i, j))
            i = j
        else:
            i += 1

    for s, e in runs:
        length = e - s
        all_table = all(blocks[k].kind == "table" for k in range(s, e))
        near_heading = any(heading_at[k] for k in range(max(0, s - 3), s))
        if near_heading:
            need = 2
        elif all_table:
            need = 3
        else:
            need = min_run
        if length >= need:
            for k in range(s, e):
                acr, exp = pairs[k]
                doc.glossary_defs.setdefault(acr, exp)
    return doc


# ----------------------------------------------------- built-in acronym reference
# A standing, document-independent acronym dictionary (the organization's "house"
# list -- e.g. the New Hire Acronym List), shipped as reference_acronyms.csv next to
# this module. Any reference acronym that actually APPEARS in a document is treated as
# a known/"existing" acronym with the house expansion -- so the tool recognizes
# org-standard acronyms a document never bothers to define, without flooding the report
# with the whole dictionary. The document's own definition always wins (setdefault), and
# a document that defines an acronym differently still surfaces as a conflict downstream.
_REF_CACHE = None


def load_reference_acronyms(path=None):
    """Load the bundled acronym reference (an `acronym,definition` CSV) -> dict, cached.
    Returns {} if the file is absent, so the pipeline runs fine without it."""
    global _REF_CACHE
    if path is None and _REF_CACHE is not None:
        return _REF_CACHE
    import csv
    p = Path(path) if path else (Path(__file__).resolve().parent / "reference_acronyms.csv")
    out = {}
    if p.exists():
        with p.open(encoding="utf-8-sig", newline="") as fh:
            for row in csv.reader(fh):
                if len(row) < 2:
                    continue
                acr, exp = row[0].strip(), row[1].strip()
                if not acr or not exp or acr.lower() in ("acronym", "abbreviation"):
                    continue                       # skip header / blank rows
                out.setdefault(acr, exp)
    if path is None:
        _REF_CACHE = out
    return out


def apply_reference_acronyms(doc, ref=None):
    """Merge reference acronyms that appear in the document into doc.glossary_defs
    (the document's own definitions take precedence). Returns the count merged."""
    ref = load_reference_acronyms() if ref is None else ref
    if not ref:
        return 0
    keys = sorted(ref, key=len, reverse=True)      # longest-first so MIL-STD beats MIL
    pat = re.compile(r"(?<![A-Za-z0-9])(?:" + "|".join(re.escape(k) for k in keys)
                     + r")(?![A-Za-z0-9])")         # case-sensitive, boundary-anchored
    n = 0
    for acr in set(pat.findall(doc.text)):
        if acr not in doc.glossary_defs:
            doc.glossary_defs[acr] = ref[acr]
            n += 1
    return n


def load(path) -> Doc:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return from_docx(path)
    if ext == ".xml":
        return from_xml(path)
    if ext == ".pdf":
        return from_pdf(path)
    raise ValueError(f"Unsupported file type: {ext} (expected .docx, .pdf, or .xml)")
