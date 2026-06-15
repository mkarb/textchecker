import importlib, sys, os, zipfile, shutil, tempfile
import extract, checks
importlib.reload(extract); importlib.reload(checks)
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

TMP = tempfile.gettempdir()
def tp(name): return os.path.join(TMP, name)

P, F = [], []
def ok(n, c): (P if c else F).append(n); print(("PASS " if c else "FAIL ") + n)

def field(p, instr, cached):
    r = OxmlElement('w:r'); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'),'begin'); r.append(fc); p._p.append(r)
    r = OxmlElement('w:r'); it = OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text = f" {instr} "; r.append(it); p._p.append(r)
    r = OxmlElement('w:r'); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'),'separate'); r.append(fc); p._p.append(r)
    if cached is not None:
        r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = str(cached); r.append(t); p._p.append(r)
    r = OxmlElement('w:r'); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'),'end'); r.append(fc); p._p.append(r)

def page_field_para(p):
    p.add_run("Page ")
    field(p, "PAGE", 3)
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.set(qn('xml:space'),'preserve'); t.text = " of "; r.append(t); p._p.append(r)
    fs = OxmlElement('w:fldSimple'); fs.set(qn('w:instr'), " NUMPAGES ")
    rr = OxmlElement('w:r'); tt = OxmlElement('w:t'); tt.text = "42"; rr.append(tt); fs.append(rr); p._p.append(fs)

# ---------- 0) the python-docx deficiency, on a standalone doc ----------
d0 = Document(); page_field_para(d0.sections[0].footer.paragraphs[0])
ok("python-docx .text drops fldSimple cached text", d0.sections[0].footer.paragraphs[0].text == "Page 3 of ")
ok("_w_text recovers full field text", extract._w_text(d0.sections[0].footer.paragraphs[0]._p) == "Page 3 of 42")

# ---------- 1) sections fixture: add sections FIRST, then write footers ----------
d = Document()
d.add_paragraph("The widget assembly is described in this manual.")
d.add_paragraph("Details are listed in the appendix, see page 99 for the full matrix.")
d.add_paragraph("Torque values are given per page 7 of this manual.")
d.add_section(); d.add_section()
s0, s1, s2 = d.sections[0], d.sections[1], d.sections[2]
page_field_para(s0.footer.paragraphs[0])                      # field-based Page X of Y
s1.footer.is_linked_to_previous = False
s1.footer.paragraphs[0].text = "ACME Proprietary Page 9 of 42" # different cached page
# s2 stays linked -> must not duplicate
d.save(tp("cov_sections.docx"))

doc = extract.load(tp("cov_sections.docx"))
pages = [b for b in doc.blocks if b.kind == "page"]
footers = [b for b in doc.blocks if b.kind == "footer"]
ok("ONE synthetic pagination block", len(pages) == 1 and pages[0].text == "Page 1 of 42")
ok("footer residual kept, page marker stripped",
   any("ACME Proprietary" in b.text and "Page 9" not in b.text for b in footers))
ok("linked section adds no duplicate footer", sum("ACME" in b.text for b in footers) == 1)

findings, toc, total = checks.page_checks(doc)
probs = [x.problem for x in findings]
ok("page total detected = 42", total == 42)
ok("no fake page_gap/out_of_order/duplicate from section footers",
   not any(p in probs for p in ("page_gap", "out_of_order_page", "duplicate_page_number")))
ok("xref out-of-range fires (page 99 > 42)",
   any(x.problem == "xref_out_of_range" and "99" in x.detail for x in findings))
ok("in-range xref (page 7) not flagged", not any("page 7 " in x.detail for x in findings))

# ---------- 2) footnotes via package surgery ----------
FOOT = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
 '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
 '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
 '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
 '<w:footnote w:id="2"><w:p><w:r><w:t>Recieve inspection occurs at the dock.</w:t></w:r></w:p></w:footnote>'
 '</w:footnotes>')
with zipfile.ZipFile(tp("cov_sections.docx")) as zin, \
     zipfile.ZipFile(tp("cov_full.docx"), "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.namelist():
        data = zin.read(item)
        if item == "[Content_Types].xml":
            data = data.replace(b"</Types>",
                b'<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/></Types>')
        if item == "word/_rels/document.xml.rels":
            data = data.replace(b"</Relationships>",
                b'<Relationship Id="rIdFn9" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/></Relationships>')
        zout.writestr(item, data)
    zout.writestr("word/footnotes.xml", FOOT)
doc2 = extract.load(tp("cov_full.docx"))
fns = [b for b in doc2.blocks if b.kind == "footnote"]
ok("footnote extracted (separators skipped)", len(fns) == 1 and "Recieve inspection" in fns[0].text)
ok("footnote text reaches doc.text", "Recieve" in doc2.text)

# ---------- 3) textbox: Choice + Fallback, extracted once ----------
d3 = Document()
d3.add_paragraph("Body paragraph before the shape.")
NS = ('xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
      'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
      'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
      'xmlns:v="urn:schemas-microsoft-com:vml"')
AC = (f'<mc:AlternateContent {NS}>'
      '<mc:Choice Requires="wps"><wps:txbx><w:txbxContent>'
      '<w:p><w:r><w:t>Caution decal: do not exceed 42 PSI</w:t></w:r></w:p>'
      '</w:txbxContent></wps:txbx></mc:Choice>'
      '<mc:Fallback><v:textbox><w:txbxContent>'
      '<w:p><w:r><w:t>Caution decal: do not exceed 42 PSI</w:t></w:r></w:p>'
      '</w:txbxContent></v:textbox></mc:Fallback>'
      '</mc:AlternateContent>')
run = d3.add_paragraph().add_run()
run._r.append(etree.fromstring(AC))
d3.save(tp("cov_txbx.docx"))
doc3 = extract.load(tp("cov_txbx.docx"))
boxes = [b for b in doc3.blocks if b.kind == "textbox"]
ok("textbox extracted exactly once (Fallback skipped)", len(boxes) == 1 and "Caution decal" in boxes[0].text)
ok("no duplicate in doc.text", doc3.text.count("Caution decal") == 1)

# ---------- 4) nested table ----------
d4 = Document()
t = d4.add_table(rows=1, cols=2)
t.rows[0].cells[0].text = "Outer cell A"
nested = t.rows[0].cells[1].add_table(rows=1, cols=2)
nested.rows[0].cells[0].text = "NestedTerm"
nested.rows[0].cells[1].text = "Nested Definition"
d4.save(tp("cov_nested.docx"))
doc4 = extract.load(tp("cov_nested.docx"))
ok("nested table row emitted with nested path",
   any("NestedTerm | Nested Definition" in b.text and b.path.count("tbl") == 2 for b in doc4.blocks))
ok("nested text not duplicated", doc4.text.count("NestedTerm") == 1)

print("\n=== %d passed, %d failed ===" % (len(P), len(F)))
sys.exit(1 if F else 0)
