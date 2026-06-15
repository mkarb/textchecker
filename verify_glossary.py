import importlib, sys, os, tempfile
import extract, checks
importlib.reload(extract); importlib.reload(checks)
from docx import Document

TMP = tempfile.gettempdir()
STORM = os.path.join(TMP, "storm_like.docx")
GLOSSP = os.path.join(TMP, "gloss_paragraphs.docx")

# ---- fixture A: STORM II-like glossary TABLE + prose + two decoy tables ----
doc = Document()
doc.add_heading("Acronyms and Abbreviations", level=1)
t = doc.add_table(rows=0, cols=2)
for a, b in [("Acronym", "Definition"),                       # header row -> skipped
             ("CA", "Corrective Action"),
             ("FMECA", "Failure Modes, Effects, and Criticality Analysis"),
             ("PS", "Performance Specification"),
             ("DR", "Design Review"),
             ("CI", "Configuration Item"),
             ("STORM", "Small Tactical Optical Rifle Mounted"),
             ("SD-18", "NAVSEA Standard Drawing 18"),
             ("I/O", "Input Output")]:
    c = t.add_row().cells; c[0].text = a; c[1].text = b
doc.add_heading("1. Scope", level=1)
doc.add_paragraph("FMECA analysis was performed on all replaceable units in the system.")
doc.add_paragraph("Failure modes were identified and ranked by severity classification.")
bom = doc.add_table(rows=0, cols=4)                            # decoy: BOM
for r in [("Laser Diode 947nm 80W", "LAS-127-01", "ST2-025-A5", "1"),
          ("CAP 10uF 50V", "12105C106MAT2A", "ST2-611-A1", "2"),
          ("Resistor 10k 0603", "RES-104", "ST2-330-B2", "5")]:
    cc = bom.add_row().cells
    for i, v in enumerate(r): cc[i].text = v
risk = doc.add_table(rows=0, cols=2)                           # decoy: risk phrases
for r in [("Alignment compensating porro mirror Active Q-switch technology", "Mitigated"),
          ("Proven components from STORM SLX Laser performance over temperature", "Open")]:
    cc = risk.add_row().cells
    for i, v in enumerate(r): cc[i].text = v
doc.save(STORM)

# ---- fixture B: paragraph/tab-style glossary (no table) ----
doc2 = Document()
doc2.add_heading("List of Acronyms", level=1)
for a, b in [("RCM", "Reliability Centered Maintenance"),
             ("LRU", "Line Replaceable Unit"),
             ("BIT", "Built In Test")]:
    doc2.add_paragraph(a + "\t" + b)
doc2.add_heading("Body", level=1)
doc2.add_paragraph("The system uses RCM principles throughout.")
doc2.save(GLOSSP)

P, F = [], []
def ok(n, c): (P if c else F).append(n); print(("PASS " if c else "FAIL ") + n)

d = extract.load(STORM); g = d.glossary_defs
expect = {"CA": "Corrective Action",
          "FMECA": "Failure Modes, Effects, and Criticality Analysis",
          "PS": "Performance Specification", "DR": "Design Review",
          "CI": "Configuration Item", "STORM": "Small Tactical Optical Rifle Mounted",
          "SD-18": "NAVSEA Standard Drawing 18", "I/O": "Input Output"}
for k, v in expect.items():
    ok("harvested " + k, g.get(k) == v)
ok("header 'Acronym' skipped", "Acronym" not in g)
ok("prose 'Failure' not harvested", "Failure" not in g)
ok("decoy BOM not harvested", not any(k.startswith("CAP") for k in g))
ok("decoy risk-phrase not harvested", not any(" " in k for k in g))
ok("exactly the 8 expected keys", set(g) == set(expect))

findings, defs = checks.acronym_consistency(d)        # UNCHANGED pass consumes glossary_defs
ok("CA -> existing_acronyms", "Corrective Action" in defs.get("CA", []))
ok("FMECA -> existing (commas verbatim)",
   "Failure Modes, Effects, and Criticality Analysis" in defs.get("FMECA", []))
ok("SD-18 -> existing", "NAVSEA Standard Drawing 18" in defs.get("SD-18", []))

d2 = extract.load(GLOSSP)
ok("paragraph glossary RCM", d2.glossary_defs.get("RCM") == "Reliability Centered Maintenance")
ok("paragraph glossary LRU", d2.glossary_defs.get("LRU") == "Line Replaceable Unit")
ok("paragraph glossary BIT", d2.glossary_defs.get("BIT") == "Built In Test")

d3 = extract.load("sample_manual.xml")
ok("sample_manual.xml unaffected", d3.glossary_defs == {})

print("\n=== %d passed, %d failed ===" % (len(P), len(F)))
sys.exit(1 if F else 0)
